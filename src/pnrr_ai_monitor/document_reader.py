from __future__ import annotations

import logging
import warnings
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
import shutil
import subprocess
import tempfile

import docx
import pdfplumber
from bs4 import BeautifulSoup, XMLParsedAsHTMLWarning
from pypdf import PdfReader

# pdfminer (via pdfplumber) and pypdf log routine noise for genuinely-real
# but imperfect PDFs (missing FontBBox, malformed ToUnicode maps, undecodable
# characters) — not exceptions, just logging, so they print regardless of
# what our code does with the return value. Quiet them at import time rather
# than building a "quality status" tracking system nobody downstream needs.
logging.getLogger("pdfminer").setLevel(logging.ERROR)
logging.getLogger("pypdf").setLevel(logging.ERROR)
# BeautifulSoup's own encoding-detection heuristic (bs4.dammit.UnicodeDammit)
# logs "Some characters could not be decoded..." routinely for real-world
# HTML with ambiguous/mixed character encodings — same kind of harmless
# library noise as the two above, just from _read_html's parser instead of
# the PDF path, and missed by the original fix because that was only tested
# against a PDF document.
logging.getLogger("bs4").setLevel(logging.ERROR)
# Belt-and-suspenders alongside the content-sniffing fix below: if some
# future caller still hands genuinely-XML bytes to _read_html, don't let the
# warning through either.
warnings.filterwarnings("ignore", category=XMLParsedAsHTMLWarning)


@dataclass(frozen=True)
class ReadDocumentResult:
    text: str
    content_type: str
    used_ocr: bool = False


class DocumentReader:
    def read(self, content: bytes, content_type: str, url: str) -> ReadDocumentResult:
        lowered = content_type.lower()
        low_url = url.lower()
        if (low_url.endswith(".pdf") or "pdf" in lowered) and content.lstrip()[:5] == b"%PDF-":
            return self._read_pdf(content, content_type)
        # Some servers label error/redirect pages (or plain HTML) with a .pdf
        # URL or a "pdf" content-type header even though the bytes aren't a
        # PDF at all — routing those into the PDF path wastes two failed
        # library parses plus a doomed pdftoppm/OCR subprocess spawn for
        # content that was never a PDF to begin with. Fall through to HTML.
        if self._looks_like_xml(content):
            # Same principle as the PDF check above: trust the actual bytes,
            # not the content-type header or URL, which can mislabel a real
            # RSS/Atom/sitemap response (seen from search-fallback results
            # and other arbitrary URLs) as HTML.
            return self._read_xml(content, content_type)
        if low_url.endswith(".docx") or "wordprocessingml" in lowered:
            return self._read_docx(content, content_type)
        if low_url.endswith(".doc"):
            # Legacy binary Word format — python-docx (and our other tools) can't
            # read it. The real notice text is always in the signed PDF alongside
            # it, so we skip this quietly rather than mis-parsing it as HTML.
            return ReadDocumentResult(text="", content_type=content_type)
        return self._read_html(content, content_type)

    @staticmethod
    def _looks_like_xml(content: bytes) -> bool:
        head = content.lstrip()[:10].lower()
        return head.startswith((b"<?xml", b"<rss", b"<feed", b"<urlset"))

    def _read_xml(self, content: bytes, content_type: str) -> ReadDocumentResult:
        try:
            root = ET.fromstring(content)
        except ET.ParseError:
            # Malformed feed — fall back rather than lose the content
            # entirely; ElementTree is stricter than BeautifulSoup and will
            # raise on real-world XML that isn't perfectly well-formed.
            return self._read_html(content, content_type)
        parts = [elem.text.strip() for elem in root.iter() if elem.text and elem.text.strip()]
        text = " ".join(" ".join(parts).split())
        return ReadDocumentResult(text=text, content_type=content_type)

    def _read_docx(self, content: bytes, content_type: str) -> ReadDocumentResult:
        try:
            document = docx.Document(BytesIO(content))
            parts = [p.text for p in document.paragraphs if p.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text)
            text = " ".join(" ".join(parts).split())
            return ReadDocumentResult(text=text, content_type=content_type)
        except Exception:
            return ReadDocumentResult(text="", content_type=content_type)

    def _read_html(self, content: bytes, content_type: str) -> ReadDocumentResult:
        soup = BeautifulSoup(content, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = " ".join(soup.get_text(" ", strip=True).split())
        return ReadDocumentResult(text=text, content_type=content_type)

    # Read the whole notice, not just the first pages: the CUP/CLP and the
    # "esperto esterno" clause frequently appear late in the document. Cap only
    # as a safety value against pathological files.
    MAX_PDF_PAGES = 60
    MAX_OCR_PAGES = 12
    # Above this fraction of U+FFFD (REPLACEMENT CHARACTER) in the extracted text,
    # treat the extraction as unreliable (broken/embedded font map) rather than
    # silently trusting a possibly-garbled read.
    MAX_REPLACEMENT_RATIO = 0.02

    def _read_pdf(self, content: bytes, content_type: str) -> ReadDocumentResult:
        # pdfplumber extracts these stamped/signed school PDFs reliably; pypdf
        # often returns almost nothing for them. Try pdfplumber first, fall back
        # to pypdf, then to the OCR hook for genuinely scanned (image) PDFs or
        # ones whose font mapping is too corrupted to trust either extractor.
        primary = self._extract_pdfplumber(content)
        if self._is_usable(primary):
            return ReadDocumentResult(text=primary, content_type=content_type)
        secondary = self._extract_pypdf(content)
        if self._is_usable(secondary):
            return ReadDocumentResult(text=secondary, content_type=content_type)
        return self._ocr_fallback(content, content_type)

    def _is_usable(self, text: str) -> bool:
        return len(text) >= 80 and self._replacement_ratio(text) <= self.MAX_REPLACEMENT_RATIO

    @staticmethod
    def _replacement_ratio(text: str) -> float:
        if not text:
            return 0.0
        return text.count("�") / len(text)

    def _extract_pdfplumber(self, content: bytes) -> str:
        try:
            with pdfplumber.open(BytesIO(content)) as pdf:
                parts = [(p.extract_text() or "") for p in pdf.pages[: self.MAX_PDF_PAGES]]
            return " ".join(" ".join(parts).split())
        except Exception:
            return ""

    def _extract_pypdf(self, content: bytes) -> str:
        try:
            reader = PdfReader(BytesIO(content))
            parts = [(page.extract_text() or "") for page in reader.pages[: self.MAX_PDF_PAGES]]
            return " ".join(" ".join(parts).split())
        except Exception:
            return ""

    def _ocr_fallback(self, content: bytes, content_type: str) -> ReadDocumentResult:
        # OCR is intentionally optional: if Poppler/Tesseract are not installed,
        # the monitor keeps running and simply treats this scanned PDF as unread.
        if not shutil.which("pdftoppm") or not shutil.which("tesseract"):
            return ReadDocumentResult(text="", content_type=content_type, used_ocr=True)
        try:
            with tempfile.TemporaryDirectory(prefix="pnrr-ocr-") as tmp:
                work = Path(tmp)
                pdf_path = work / "document.pdf"
                pdf_path.write_bytes(content)
                prefix = work / "page"
                subprocess.run(
                    [
                        "pdftoppm",
                        "-f", "1",
                        "-l", str(self.MAX_OCR_PAGES),
                        "-r", "200",
                        "-png",
                        str(pdf_path),
                        str(prefix),
                    ],
                    check=True,
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                )
                parts: list[str] = []
                for image in sorted(work.glob("page-*.png")):
                    text = self._run_tesseract(image, "ita+eng") or self._run_tesseract(image, "eng")
                    if text:
                        parts.append(text)
                return ReadDocumentResult(text=" ".join(" ".join(parts).split()), content_type=content_type, used_ocr=True)
        except Exception:
            return ReadDocumentResult(text="", content_type=content_type, used_ocr=True)

    @staticmethod
    def _run_tesseract(image_path: Path, lang: str) -> str:
        try:
            result = subprocess.run(
                ["tesseract", str(image_path), "stdout", "-l", lang],
                check=True,
                capture_output=True,
                text=True,
                timeout=45,
            )
            return result.stdout or ""
        except Exception:
            return ""
